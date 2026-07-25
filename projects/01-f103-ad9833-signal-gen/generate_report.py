#!/usr/bin/env python3
"""
生成 AD9833 信号源项目报告 — DOCX 格式
创新创业工程与实践课程
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

def add_heading_styled(text, level=1):
    """添加居中标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 0:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def add_para(text, bold=False, indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)  # 两字符缩进
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    run.bold = bold
    return p

def add_code(text):
    """添加代码块（等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    return p

def add_bullet(text):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    return p

def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 灰色背景
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3',
            qn('w:val'): 'clear'
        })
        shading.append(shd)
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # 表后空行
    return table


# ============================================================
# 封面区域（不另起一页，直接放标题区）
# ============================================================
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('基于 STM32 与 AD9833 的\n可编程多功能信号源设计')
run.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('—— 全国大学生电子设计大赛 仪表方向备赛项目')
run.font.name = '楷体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    '课程名称：创新创业工程与实践',
    '项目类别：电子设计竞赛 — 信号源模块',
    '开发平台：STM32F103C8T6 + AD9833 DDS 模块',
    '完成日期：2026 年 7 月',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 摘要
# ============================================================
add_heading_styled('摘  要', level=1)

add_para(
    '本报告详细记录了一个基于 STM32F103C8T6 微控制器与 AD9833 DDS（直接数字频率合成）芯片的可编程信号源'
    '项目的完整设计、实现与调试过程。该信号源能够通过串口命令实时控制输出信号的频率与波形，支持正弦波、三角'
    '波和方波三种波形，频率范围为 0.1 Hz 至 12.5 MHz，频率分辨率达到 0.093 Hz。本项目是全国大学生电子设计大赛'
    '（以下简称"电赛"）仪表方向各赛的核心准备工作之一，旨在为后续的信号测量、锁相放大、阻抗分析等仪器类'
    '题目提供可控的信号激励源。'
)

add_para(
    '在项目实施过程中，本文作者严格遵循"数据手册驱动开发"的方法论——从 AD9833 芯片数据手册的时序图和寄存器'
    '表出发，推导出 SPI 通信参数（CPOL=1、CPHA=0，即 Mode 2）、每一位控制位的宏定义、频率字计算公式'
    '（freq_word = f_out × 2^28 / 25MHz），并最终在 STM32CubeIDE 中完成 CubeMX 外设配置、HAL 库驱动程序编写和'
    '串口命令交互功能的实现。经过示波器实测，三种波形的频率精度和波形质量均符合预期。开发过程中发现并修复了'
    '两个关键技术问题——SPI 时钟相位配置错误与波形切换时寄存器位覆盖 bug，深刻认识到"不凭感觉、严格对照手册"'
    '在嵌入式开发中的重要性。'
)

add_para(
    '**关键词**：AD9833；DDS 信号源；STM32F103C8T6；SPI 通信；电子设计大赛；仪表方向'
)

# ============================================================
# 1. 项目背景与意义
# ============================================================
add_heading_styled('一、项目背景与意义', level=1)

add_heading_styled('1.1 全国大学生电子设计大赛简介', level=2)

add_para(
    '全国大学生电子设计大赛是教育部与工业和信息化部共同发起的大学生学科竞赛，每两年举办一届，是我国电子信息类'
    '领域规模最大、影响力最广的本科生赛事之一。比赛采用"半封闭、相对集中"的组织方式，参赛队伍需在四天三夜的'
    '时间内，根据赛题要求完成从方案设计、电路制作、程序编写到报告撰写的完整流程。电赛题目通常分为电源类、信号'
    '类、控制类、仪器仪表类、通信类等多个方向，其中仪器仪表方向（简称"仪表方向"）重点考察参赛选手对微弱信号'
    '检测、精密测量、传感器信号调理等技术的掌握程度。'
)

add_heading_styled('1.2 信号源在电赛仪表题中的核心地位', level=2)

add_para(
    '在仪表方向的赛题中，无论是要搭建锁相放大器（LIA）检测微弱信号、做阻抗分析仪测量电容/电感参数，还是设计'
    '功率计进行电能质量分析，都需要一个高精度、可编程的信号源作为激励或参考。传统的商用信号发生器价格昂贵'
    '（数千至上万元）、体积大，不符合电赛"低成本、可自制"的要求。因此，独立设计和制作一个基于 DDS 技术的'
    '信号源模块，不仅满足电赛对自制设备的要求，还能在赛前积累关键的嵌入式系统开发经验。'
)

add_heading_styled('1.3 DDS 技术简介', level=2)

add_para(
    'DDS（Direct Digital Synthesis，直接数字频率合成）是一种从相位概念出发直接合成所需波形的频率合成技术。'
    '其核心原理是：将一个周期波形的形状数字化后存入 ROM（只读存储器）中，通过相位累加器在每个时钟周期不断'
    '累加一个频率控制字，然后用累加后的相位值作为地址去查表，再经 DAC（数模转换器）和低通滤波器还原为模拟波形。'
    'DDS 相较于传统 PLL（锁相环）频率合成，具有频率分辨率极高（可达 μHz 级）、频率切换速度快（ns 级）、相位'
    '噪声低等显著优势。AD9833 是 Analog Devices 公司推出的低功耗、可编程 DDS 波形发生器芯片，内部集成了相位'
    '累加器、正弦查找表 ROM 和 10 位 DAC，仅需一个外部时钟源和少量被动元件即可输出高质量的正弦波、三角波和方波，'
    '是嵌入式信号源设计的理想选择。'
)

# ============================================================
# 2. 系统方案设计
# ============================================================
add_heading_styled('二、系统总体方案设计', level=1)

add_heading_styled('2.1 系统架构', level=2)

add_para(
    '本信号源系统采用"MCU + DDS 芯片"的主从结构，系统框架如图 1 所示。上位机（串口调试助手）通过 UART 发送'
    '命令给 STM32 主控，STM32 解析命令后通过 SPI 总线控制 AD9833 DDS 模块，AD9833 根据指令输出相应频率和波形'
    '的模拟信号。示波器用于观察和验证输出波形。'
)

add_code(
    '┌──────────┐   UART    ┌───────────┐   SPI     ┌──────────┐   模拟信号   ┌──────┐\n'
    '│ 串口助手  │──────────▶│ STM32F103 │─────────▶│ AD9833   │───────────▶│ 示波器 │\n'
    '│  (上位机) │           │  (主控)    │          │  (DDS)   │            │      │\n'
    '└──────────┘           └───────────┘          └──────────┘            └──────┘\n'
    '                           │                      │\n'
    '                      USB供电              25MHz 时钟\n'
    '                                              ↓\n'
    '                                    0.1Hz ~ 12.5MHz 信号输出'
)

add_heading_styled('2.2 硬件选型', level=2)

add_table(
    ['模块', '型号/参数', '选型理由'],
    [
        ['主控 MCU', 'STM32F103C8T6 (Blue Pill)', 'ARM Cortex-M3, 72MHz, SPI+UART丰富, 参考资料多, 价格低'],
        ['DDS 模块', 'AD9833 V40 模块', '集成 25MHz 有源晶振, 3.3V 供电, 芯片手册详尽, 模块即插即用'],
        ['调试工具', 'USB-TTL 串口模块', '3.3V 电平匹配 STM32, 波特率 115200'],
        ['测量设备', '数字示波器', '验证输出频率和波形'],
        ['开发环境', 'STM32CubeIDE', 'HAL 库生成方便, CubeMX 图形化配置外设'],
    ]
)

add_heading_styled('2.3 通信接口设计', level=2)

add_para(
    '本项目涉及两种通信接口：一是STM32与上位机之间的UART串口通信（用于人机交互），二是STM32与AD9833之间的'
    'SPI通信（用于芯片控制）。两者均通过 STM32 的硬件外设实现。UART 配置为 115200/8/N/1 并开启接收中断，SPI 配'
    '置为 Mode 2（CPOL=1, CPHA=0）、16-bit 数据帧、MSB 优先、波特率分频 256（约 281 kHz，远低于 AD9833 的'
    '40 MHz 上限）。需要特别说明的是，AD9833 要求每发送 16-bit 命令后必须将 FSYNC 信号拉高一次以锁存命令，这与'
    '普通 SPI 设备可在片选有效期间连续发送多字节不同，因此本项目放弃硬件 NSS 片选信号，改用 PA4 引脚软件模拟 FSYNC。'
)

# ============================================================
# 3. 硬件设计
# ============================================================
add_heading_styled('三、硬件电路设计', level=1)

add_heading_styled('3.1 STM32 时钟配置', level=2)

add_para(
    'STM32 外部 8MHz 晶振作为 HSE 时钟源，经 PLL 9 倍频至 72MHz 作为系统时钟（SYSCLK）。APB2 总线（72MHz）'
    '承载 SPI1 外设，APB1 总线（36MHz）承载 USART1 外设。该配置是 STM32F103C8T6 的最优性能配置。'
)

add_heading_styled('3.2 AD9833 模块电路分析', level=2)

add_para(
    'AD9833 V40 模块采用 AD9833BRMZ（10脚 MSOP 封装）芯片，板载 25MHz 有源晶振提供主时钟（MCLK）。模块引出'
    'FSYNC（帧同步，低电平有效）、SCLK（SPI 时钟）、SDATA（SPI 数据输入）三个通信引脚，以及 VCC 和 GND 供'
    '电引脚。模块输出引脚 VOUT 直接连接到芯片 DAC 输出端，无后级运放缓冲，输出幅度约为 0.6Vpp。模块支持 2.3V 至'
    '5.5V 宽电压供电，本项目中采用 3.3V 与 STM32 统一供电。'
)

add_heading_styled('3.3 接线方案', level=2)

add_table(
    ['STM32 引脚', '功能', '连接至', 'AD9833 引脚', '线色建议'],
    [
        ['PA4', 'GPIO Output (FSYNC)', '────▶', 'FSYNC', '黄色'],
        ['PA5', 'SPI1_SCK', '────▶', 'SCLK', '绿色'],
        ['PA7', 'SPI1_MOSI', '────▶', 'SDATA', '蓝色'],
        ['3.3V', '电源正', '────▶', 'VCC', '红色'],
        ['GND', '电源地', '────▶', 'GND', '黑色'],
    ]
)

add_para(
    '整个项目仅需 5 根杜邦线即可完成硬件连接（不包含调试用的串口线），体现了"最小化布线、最大化可靠性"的硬件设计原则。'
)

# ============================================================
# 4. 软件设计
# ============================================================
add_heading_styled('四、软件设计', level=1)

add_heading_styled('4.1 软件架构总览', level=2)

add_para(
    '软件包括三层：底层驱动层（ad9833.h + ad9833.c），负责 SPI 寄存器级操作和位运算；中间层（'  # wrapped
    'AD9833_WriteRegister / SetFreq / SetWave 函数），将数据手册中的操作流程转化为可调用的函数接口；'
    '应用层（main.c），通过 UART 中断接收用户命令并调用驱动函数。'
)

add_heading_styled('4.2 寄存器宏定义 — "积木式"设计', level=2)

add_para(
    '本项目独创性地采用"积木式"宏定义策略，将 AD9833 控制寄存器（16-bit）的每一位定义为独立的宏常量。'
    '这种设计的优点在于：每一个宏定义都能在数据手册中找到明确对应的 bit 位（如 BIT_B28 对应 D13、BIT_MODE 对'
    '应 D1），保证代码与手册之间的可追溯性；波形组合通过按位或（|）操作实现，直观且不易出错。'
)

add_code(
    '/* ── 第一组: 命令类型 (D15-D14) ── */\n'
    '#define CMD_WRITE_CTRL   0x0000  // 00 = 写控制寄存器\n'
    '#define CMD_WRITE_FREQ0  0x4000  // 01 = 写频率寄存器0\n'
    '#define CMD_WRITE_FREQ1  0x8000  // 10 = 写频率寄存器1\n'
    '#define CMD_WRITE_PHASE0 0xC000  // 11 = 写相位寄存器0\n'
    '\n'
    '/* ── 第二组: 控制寄存器各位 (D13-D0) ── */\n'
    '#define BIT_B28      ((uint16_t)1 << 13)  // 连续 28-bit 写频率\n'
    '#define BIT_RESET    ((uint16_t)1 << 8)   // 复位（配参数时要置 1）\n'
    '#define BIT_OPBITEN  ((uint16_t)1 << 5)   // 方波输出使能\n'
    '#define BIT_DIV2     ((uint16_t)1 << 3)   // 方波二分频\n'
    '#define BIT_MODE     ((uint16_t)1 << 1)   // 0=正弦, 1=三角\n'
    '\n'
    '/* ── 第三组: 波形组合 ── */\n'
    '#define CTRL_SINE     (BIT_RESET | BIT_B28)\n'
    '#define CTRL_TRIANGLE (BIT_RESET | BIT_B28 | BIT_MODE)\n'
    '#define CTRL_SQUARE   (BIT_RESET | BIT_B28 | BIT_OPBITEN | BIT_DIV2)'
)

add_heading_styled('4.3 核心驱动函数', level=2)

add_para(
    '驱动层提供三个核心函数，分别对应初始化、设频率、设波形三种操作：'
)

add_para(
    '（1）AD9833_WriteRegister(uint16_t data)：最底层的 SPI 写操作。按照 AD9833 数据手册第 7 页的时序图，'
    '先拉低 FSYNC → 调用 HAL_SPI_Transmit 发送 16-bit 数据 → 拉高 FSYNC 锁存命令。该函数是所有上层操作的"基石"。'
)

add_para(
    '（2）AD9833_SetFreq(float freq_hz)：将目标频率值代入数据手册第 8 页给出的公式——频率字 = f_out × 2²⁸ '
    '/ 25MHz，算出 28-bit 频率字，再拆分为低 14 位和高 14 位，通过连续两次 FREQ0 写命令发送。B28=1 模式下，'
    'AD9833 自动拼接两次写入的数据为完整的 28-bit 频率字。'
)

add_para(
    '（3）AD9833_SetWave(uint16_t wave_ctrl)：切换输出波形。先将新波形控制字（含 RESET=1）写入芯片，使 AD9833 '
    '在切换过程中保持复位状态；再写入"wave_ctrl & ~BIT_RESET"（位掩码清零 RESET 位，其余位保留）释放复位。'
    '这一"先复位→改配置→释放复位"的流程源于数据手册第 9 页的操作建议，能有效避免波形切换瞬间产生毛刺。'
)

add_heading_styled('4.4 串口命令交互', level=2)

add_para(
    '为实现灵活的人机交互，应用层实现了基于 UART 中断接收的串口命令解析功能。系统支持的命令包括：'
)

add_table(
    ['命令', '格式示例', '功能说明'],
    [
        ['freq', 'freq 1000', '设置输出频率为 1000 Hz'],
        ['wave', 'wave 0 / wave 1 / wave 2', '切换波形：0=正弦, 1=三角, 2=方波'],
        ['sweep', 'sweep 100 10000 100 50', '扫频：100Hz→10kHz, 步进100Hz, 每步50ms'],
        ['status', 'status', '查询当前频率和波形状态'],
        ['help', 'help', '显示所有可用命令'],
    ]
)

add_para(
    '命令接收采用中断模式而非轮询模式：USART1 每收到一个字符便触发中断，进入 HAL_UART_RxCpltCallback() 回调'
    '函数将字符存入环形缓冲区，检测到回车换行符（\\r\\n）后置位 rx_ready 标志。主循环检测到此标志后调用'
    'Cmd_Process() 进行命令解析与执行。这种设计避免了主循环中的忙等待，CPU 在无命令时几乎处于空闲状态。'
)

# ============================================================
# 5. 调试与问题解决
# ============================================================
add_heading_styled('五、调试过程与关键问题解决', level=1)

add_heading_styled('5.1 测试方案', level=2)

add_para(
    '测试分为三个阶段：基本功能测试（验证三种波形输出是否正常）、频率精度测试（用示波器测量不同设定频率下的'
    '实际输出频率）、串口命令测试（验证所有命令的响应正确性）。测试工具包括数字示波器（测波形/频率）、逻辑分析仪'
    '（验证 SPI 时序）和串口调试助手（发送命令）。'
)

add_heading_styled('5.2 关键问题一：SPI 模式配置错误', level=2)

add_para(
    '首次烧录程序后，用示波器观察到的波形周期接近 10 秒，而设定频率为 1kHz（周期应为 1ms）。检查代码逻辑无误'
    '后，发现问题出在 STM32CubeMX 的 SPI 配置上：CPHA 参数被错误地设为"2 Edge"（对应 CPHA=1，即 SPI Mode 3）。'
)

add_para(
    '根因分析：CubeMX 中"CPHA=2 Edge"的含义是"在第二个时钟边沿采样数据"，而 AD9833 数据手册第 7 页的时序'
    '图明确标注"Data is clocked into the AD9833 on the falling edge of SCLK"。在 CPOL=1（时钟空闲为高电平）的'
    '前提下，下降沿是第一个时钟边沿，因此 CPHA 应为 0（对应 CubeMX 中的"1 Edge"），即 SPI Mode 2。'
    '将 CPHA 改为"1 Edge"后，输出频率恢复正常。'
)

add_para(
    '反思：SPI 模式的选择不能凭直觉（"下降沿就是第二个边沿"），必须严格对照数据手册的时序图，'
    '确认 CPOL（空闲电平）和 CPHA（采样边沿序号）两个参数。'
)

add_heading_styled('5.3 关键问题二：波形切换时寄存器覆盖', level=2)

add_para(
    '频率调节功能正常工作后，发现波形切换命令无响应——无论发送 wave 0、1 还是 2，输出始终是正弦波。追踪代码后'
    '发现，AD9833_SetWave() 函数在写入波形控制字（如 CTRL_SQUARE=0x2128）后，用 BIT_B28（0x2000）来释放复位。'
    '直接写入 0x2000 意味着将整个控制寄存器的低 8 位全部清零，MODE 位、OPBITEN 位和 DIV2 位都被覆盖为 0——而'
    '正弦波恰好不需要这些位（全 0），所以"正弦波能用"完全是一个巧合。'
)

add_para(
    '修复方法：将释放复位的操作从"直接写 BIT_B28（0x2000）"改为"写 wave_ctrl & ~BIT_RESET"——用位掩码运算'
    '只翻转 RESET 位（bit 8），其余位完全保留不变。修复后三种波形均切换正常。'
)

add_para(
    '反思：修改寄存器中的某一个 bit 时，必须使用位掩码操作（&= ~BIT 清零，|= BIT 置位），绝不能写一个"看上去'
    '差不多"的全新值。测试时必须覆盖所有状态组合，不能因为一种情况 pass 就推定全对。'
)

# ============================================================
# 6. 创新点与收获
# ============================================================
add_heading_styled('六、创新点与收获体会', level=1)

add_heading_styled('6.1 创新点', level=2)

add_bullet(
    '"积木式"宏定义方法：将芯片控制寄存器的每一位独立定义为宏常量，波形组合用 OR 操作实现，使代码与数据手册'
    '之间建立了一一对应的追溯关系。该方法可复用到后续任何芯片的驱动开发中。'
)
add_bullet(
    '"数据手册驱动开发"方法论：从时序图到 CubeMX 参数，从寄存器表到宏定义，从操作流程到函数实现，每一步都有'
    '手册依据。该方法论已在备赛文档中系统化总结，成为后续项目的标准开发范式。'
)
add_bullet(
    '串口实时控制 + 扫频功能：超越了 AD9833 基本使用，加入了实时命令交互和自动扫频功能，使信号源具备了初步的'
    '"仪器化"特征，可直接用于后续的电路测试和频率响应测量。'
)

add_heading_styled('6.2 技术收获', level=2)

add_bullet('深入理解了 SPI 通信协议的四种工作模式及其与芯片时序图的对应关系。')
add_bullet('掌握了 DDS 频率合成的基本原理，能够根据频率分辨率反推频率字的计算公式。')
add_bullet('学会了 STM32CubeIDE 从 Pinout 配置 → Clock 配置 → 外设参数的全流程使用方法。')
add_bullet('体会到位运算在寄存器级编程中的核心作用，掌握了安全的位掩码操作规范。')
add_bullet('认识到"先复位→改参数→释放复位"这一通用操作流程在芯片初始化中的重要性。')

add_heading_styled('6.3 工程素养提升', level=2)

add_bullet('培养了阅读英文数据手册的能力，学会从时序图中提取通信参数、从寄存器表中推导代码结构。')
add_bullet('建立了"测试驱动验证"的思维：代码写完不是终点，用示波器/逻辑分析仪验证才算完成。')
add_bullet('形成了 Git 版本管理 + 文档同步的习惯：项目代码、开发记录、错误总结全部入库管理。')
add_bullet('理解到嵌入式开发中"硬件理解"先于"软件编写"的工程逻辑：在写一行代码前，原理图和数据手册已经决定了方案。')

# ============================================================
# 7. 结论
# ============================================================
add_heading_styled('七、结论与展望', level=1)

add_para(
    '本报告完整记录了一个基于 STM32F103C8T6 与 AD9833 的可编程信号源项目的设计全过程，从硬件选型、电路分析、'
    '通信协议理解、CubeMX 外设配置、驱动代码编写到系统调试与优化，形成了完整的闭环。该系统实现了 0.1 Hz 至 12.5 '
    'MHz 频率范围内正弦波、三角波和方波的程序控制输出，经示波器测试验证，频率精度和波形质量满足电赛仪表方向备赛需求。'
)

add_para(
    '通过本项目，作者不仅掌握了 SPI 通信、DDS 原理、HAL 库编程等具体技术技能，更重要的是形成了一套可复用的'
    '"数据手册驱动开发"方法论——拿到任何一颗新芯片，先看时序图确定通信参数，再看寄存器表构建宏定义块，最后参'
    '照操作流程编写驱动函数。该方法论已在 AD835 乘法器模块的学习中继续应用并持续完善。'
)

add_para(
    '后续工作展望：一是将本信号源与 AD835 模拟乘法器组合，搭建锁相放大器（LIA）原型，实现微弱信号的相敏检测；'
    '二是增加 LCD 显示屏和旋转编码器，脱离串口实现独立的人机交互界面；三是探索将 STM32 的 DAC 或外部高速 DAC '
    '与 FPGA 结合，实现更高频率、更高分辨率的信号源。'
)

# ============================================================
# 参考文献
# ============================================================
add_heading_styled('参考文献', level=1)

refs = [
    '[1] Analog Devices. AD9833: Low Power, 12.65 mW, 2.3 V to 5.5 V, Programmable Waveform Generator Data Sheet (Rev. E) [Z]. 2012.',
    '[2] STMicroelectronics. STM32F103x8/STM32F103xB Datasheet: Medium-density performance line ARM®-based 32-bit MCU (Rev. 19) [Z]. 2023.',
    '[3] STMicroelectronics. RM0008: STM32F101xx/STM32F102xx/STM32F103xx Reference Manual (Rev. 21) [Z]. 2021.',
    '[4] 全国大学生电子设计竞赛组委会. 全国大学生电子设计竞赛章程 [Z]. 2023.',
    '[5] 童诗白, 华成英. 模拟电子技术基础(第5版)[M]. 北京: 高等教育出版社, 2015.',
    '[6] 刘火良, 杨森. STM32库开发实战指南(第2版)[M]. 北京: 机械工业出版社, 2020.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(ref)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

# ============================================================
# 附录：实物连线与测试照片说明
# ============================================================
doc.add_page_break()
add_heading_styled('附  录', level=1)
add_heading_styled('附录 A  硬件接线表', level=2)

add_table(
    ['STM32 Blue Pill', '连接线', 'AD9833 V40 模块', '备注'],
    [
        ['PA4 (GPIO Output)', '杜邦线(黄)', 'FSYNC', '软件控制帧同步'],
        ['PA5 (SPI1_SCK)', '杜邦线(绿)', 'SCLK', 'SPI 时钟'],
        ['PA7 (SPI1_MOSI)', '杜邦线(蓝)', 'SDATA', 'SPI 数据(主出从入)'],
        ['3.3V', '杜邦线(红)', 'VCC', '3.3V 统一供电'],
        ['GND', '杜邦线(黑)', 'GND', '共地'],
        ['PA9 (USART1_TX)', 'USB-TTL(白)', 'RX', '串口发送(调试)'],
        ['PA10 (USART1_RX)', 'USB-TTL(绿)', 'TX', '串口接收(调试)'],
        ['GND', 'USB-TTL(黑)', 'GND', '串口共地'],
    ]
)

add_heading_styled('附录 B  关键参数速查', level=2)

add_table(
    ['参数项', '数值', '来源/依据'],
    [
        ['SPI 模式', 'Mode 2 (CPOL=1, CPHA=0)', 'AD9833 手册第7页时序图'],
        ['SPI 数据帧', '16-bit, MSB First', 'AD9833 手册第6页'],
        ['SPI 时钟频率', '~281 kHz', '72MHz / 256 ≈ 281kHz << 40MHz上限'],
        ['频率分辨率', '0.093 Hz', '25MHz / 2^28 ≈ 0.093Hz'],
        ['输出频率范围', '0.1 Hz ~ 12.5 MHz', '理论 Nyquist: MCLK/2 = 12.5MHz'],
        ['输出幅度', '~0.6 Vpp', 'AD9833 模块无后级运放'],
        ['频率字公式', 'f_word = f_out × 2^28 / 25MHz', 'AD9833 手册第8页'],
    ]
)

# ============================================================
# 保存
# ============================================================
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'AD9833信号源项目报告.docx')
doc.save(output_path)
print(f'报告已生成: {output_path}')
